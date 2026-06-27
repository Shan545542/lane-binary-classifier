from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "智能驾驶车道二分类实验报告.docx"
REPORT_ASSETS = ROOT / "docs" / "report_assets"


def set_font(run, font_name: str = "宋体", size: int = 10) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)


def add_paragraph(doc: Document, text: str, first_line_indent: bool = True) -> None:
    paragraph = doc.add_paragraph()
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_font(run, "宋体", 10.5)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_font(run, "黑体", 14 if level == 1 else 12)


def add_code(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(code.strip())
    set_font(run, "Consolas", 8.5)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_font(run, "宋体", 10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        run = table.rows[0].cells[index].paragraphs[0].add_run(header)
        set_font(run, "黑体", 10)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            run = cells[index].paragraphs[0].add_run(value)
            set_font(run, "宋体", 9.5)


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.4) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"图示缺失：{image_path}", first_line_indent=False)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_font(caption_run, "宋体", 9)


def build_report() -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("智能驾驶车道区域二分类感知实验报告")
    set_font(run, "黑体", 18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于前视图像 ROI 的 in_lane / out_lane 分类模型")
    set_font(run, "宋体", 11)

    add_paragraph(
        doc,
        "说明：本文档为当前阶段实验报告正式版。实验已完成 Toy 数据集流程验证，并基于 TuSimple 开源车道线数据集完成真实数据训练、测试、TensorBoard 记录、ONNX 导出和 ONNX Runtime 推理验证。",
        first_line_indent=False,
    )

    add_heading(doc, "一、实验目的", 1)
    add_paragraph(
        doc,
        "本实验面向智能驾驶感知中的基础决策场景，设计并实现一个车道区域二分类模型。模型输入车载前视图像中的局部感兴趣区域 ROI，输出该区域属于“车道内”还是“车道外”。该任务不追求精确拟合车道线几何形状，而是验证在自动驾驶感知流程中，基于局部视觉信息快速判断车辆当前位置是否处于目标车道区域内的可行性。",
    )
    add_paragraph(
        doc,
        "实验目标包括：构建轻量级深度学习二分类模型；基于开源车道线数据集 TuSimple 设计 ROI 数据转换方案；实现训练、验证、TensorBoard 可视化、ONNX 导出和 ONNX Runtime 推理；分析模型方案的优缺点，并为后续真实数据训练和部署优化提供基础。",
    )

    add_heading(doc, "二、实验原理", 1)
    add_heading(doc, "2.1 ROI 感兴趣区域", 2)
    add_paragraph(
        doc,
        "ROI 是 Region of Interest 的缩写，中文通常称为“感兴趣区域”。在车载前视图像中，整张图像包含天空、道路、车辆、路边建筑和远处背景等大量信息，但本实验真正关心的是车辆前方与当前车道相关的道路区域。因此，先从原始图像中裁剪出车辆前方道路局部区域，可以减少无关背景干扰，降低输入尺寸，并让模型更集中地学习车道线、路面边界和车道中心偏移等视觉线索。",
    )
    add_paragraph(
        doc,
        "在 TuSimple 数据转换阶段，脚本会根据标注的车道线坐标估计当前车道中心。以当前车道中心为基准裁剪得到的 ROI 作为正样本 in_lane；向左或向右偏移裁剪得到的 ROI 作为负样本 out_lane。这样可以把原始车道线检测数据转换成适合图像分类训练的二分类数据集。",
    )

    add_heading(doc, "2.2 TuSimple 标注与二分类转换", 2)
    add_paragraph(
        doc,
        "TuSimple 原始任务是车道线检测，标注文件采用 JSON Lines 格式。每一行对应一张图像，主要包含 raw_file、lanes 和 h_samples 等字段。其中 raw_file 表示图片路径，h_samples 表示若干固定的 y 坐标，lanes 中的每一组 x 坐标表示某条车道线在这些 y 坐标上的位置。x 值为 -2 时表示该位置没有有效车道线点。",
    )
    add_paragraph(
        doc,
        "本实验的任务不是直接预测车道线坐标，而是判断局部 ROI 是否位于车道区域内。因此需要利用 TuSimple 的车道线坐标估计当前车道中心，再生成 ImageFolder 风格的二分类数据目录：train/in_lane、train/out_lane、val/in_lane、val/out_lane。",
    )
    add_code(
        doc,
        """
def estimate_lane_center(record: dict, image_width: int, y_ref: int) -> int | None:
    xs = []
    for lane in record["lanes"]:
        x = lane_x_at_y(lane, record["h_samples"], y_ref)
        if x is not None:
            xs.append(x)
    xs = sorted(set(xs))
    if len(xs) < 2:
        return None

    image_center = image_width / 2
    left = [x for x in xs if x < image_center]
    right = [x for x in xs if x > image_center]
    if not left or not right:
        return None
    return int((max(left) + min(right)) / 2)
        """,
    )

    add_heading(doc, "2.3 模型结构", 2)
    add_paragraph(
        doc,
        "本实验采用轻量级卷积神经网络 LaneBinaryNet。模型由 4 个卷积块组成，每个卷积块包含 Conv2d、BatchNorm、ReLU 和 MaxPool，用于逐层提取车道线边缘、路面纹理和局部空间结构特征。随后使用 Adaptive Average Pooling 将空间特征压缩为向量，再经过 Dropout 和 Linear 层输出一个 logit。",
    )
    add_paragraph(
        doc,
        "模型输出不是直接的类别，而是一个实数 logit。推理时使用 sigmoid 函数将 logit 转换为属于 in_lane 的概率。当概率大于等于 0.5 时判定为 in_lane，否则判定为 out_lane。",
    )
    add_code(
        doc,
        """
class LaneBinaryNet(nn.Module):
    def __init__(self, dropout: float = 0.25) -> None:
        super().__init__()
        self.features = nn.Sequential(
            ConvBlock(3, 24),
            ConvBlock(24, 48),
            ConvBlock(48, 96),
            ConvBlock(96, 128),
        )
        self.classifier = nn.Sequential(
            nn.AdaptiveAvgPool2d((1, 1)),
            nn.Flatten(),
            nn.Dropout(p=dropout),
            nn.Linear(128, 1),
        )

    def forward(self, x: torch.Tensor) -> torch.Tensor:
        logits = self.classifier(self.features(x))
        return logits.squeeze(1)
        """,
    )

    add_heading(doc, "2.4 损失函数与评价指标", 2)
    add_paragraph(
        doc,
        "训练阶段使用 BCEWithLogitsLoss 作为损失函数。该损失函数内部结合了 sigmoid 与二元交叉熵，相比手动先 sigmoid 再计算损失具有更好的数值稳定性。评价指标包括准确率 accuracy、精确率 precision、召回率 recall 和 F1 分数。其中 F1 分数综合考虑 precision 与 recall，适合在类别数量可能不完全均衡时作为主要模型选择指标。",
    )
    add_code(
        doc,
        """
criterion = nn.BCEWithLogitsLoss()

probs = torch.sigmoid(logits)
preds = (probs >= 0.5).long()
precision = tp / max(tp + fp, 1)
recall = tp / max(tp + fn, 1)
f1 = 2 * precision * recall / max(precision + recall, 1e-12)
        """,
    )

    add_heading(doc, "2.5 数据增强", 2)
    add_paragraph(
        doc,
        "为提高模型泛化能力，训练集在读取图像时加入了简单增强，包括水平翻转、亮度扰动和对比度扰动。水平翻转可以模拟左右车道位置变化；亮度与对比度增强可以模拟不同光照条件。验证集不使用增强，以保证评估结果稳定。",
    )
    add_code(
        doc,
        """
if random.random() < 0.5:
    image = image.transpose(Image.FLIP_LEFT_RIGHT)
if random.random() < 0.8:
    factor = random.uniform(0.65, 1.35)
    image = ImageEnhance.Brightness(image).enhance(factor)
if random.random() < 0.3:
    factor = random.uniform(0.8, 1.2)
    image = ImageEnhance.Contrast(image).enhance(factor)
        """,
    )

    add_heading(doc, "三、实验环境与项目结构", 1)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["开发语言", "Python 3.12"],
            ["深度学习框架", "PyTorch"],
            ["图像处理", "Pillow、NumPy"],
            ["训练可视化", "TensorBoard"],
            ["模型导出", "ONNX"],
            ["部署推理验证", "ONNX Runtime"],
            ["版本管理", "Git、GitHub"],
        ],
    )
    add_paragraph(
        doc,
        "项目采用标准 GitHub 工程结构，核心代码位于 src/lane_binary_classifier，数据转换脚本位于 scripts，说明文档位于 docs，训练配置位于 configs。data、outputs、.venv、模型权重和 ONNX 文件均被 .gitignore 忽略，避免将大文件和本地环境提交到远程仓库。",
    )

    add_heading(doc, "四、实验步骤", 1)
    steps = [
        "创建独立 Python 虚拟环境，安装 requirements.txt 中的依赖，并使用 pip install -e . 将项目以可编辑模式安装。",
        "使用 make_toy_dataset.py 生成小型合成数据集，先验证训练、保存模型、TensorBoard、ONNX 导出和推理流程是否完整。",
        "实现 LaneImageFolder 数据加载器，按照 train/val 和 in_lane/out_lane 目录组织数据。",
        "实现轻量 CNN 模型 LaneBinaryNet，并使用 BCEWithLogitsLoss 训练二分类模型。",
        "训练时记录 loss、accuracy、precision、recall 和 F1，并将曲线写入 TensorBoard。",
        "根据验证集 F1 保存最优模型 best.pt，同时保存最后一轮模型 last.pt。",
        "将 best.pt 导出为 ONNX 格式，得到 lane_binary_classifier.onnx。",
        "使用 ONNX Runtime 加载 ONNX 模型，对单张 ROI 图片进行独立推理验证。",
        "使用 prepare_tusimple_binary.py 将 TuSimple 开源车道线数据转换成正式二分类 ROI 数据集，并完成真实数据训练与测试。",
    ]
    for step in steps:
        add_bullet(doc, step)

    add_heading(doc, "五、关键命令", 1)
    add_code(
        doc,
        """
python scripts\\make_toy_dataset.py --output data\\toy_lane --train-size 400 --val-size 80

python -m lane_binary_classifier.train --data-dir data\\toy_lane --output-dir outputs\\toy_run --epochs 8 --batch-size 32 --tensorboard

tensorboard --logdir outputs\\toy_run\\tensorboard

python -m lane_binary_classifier.export_onnx --checkpoint outputs\\toy_run\\best.pt --output outputs\\toy_run\\lane_binary_classifier.onnx

python -m lane_binary_classifier.predict_onnx --onnx-model outputs\\toy_run\\lane_binary_classifier.onnx --image data\\toy_lane\\val\\in_lane\\0000.png
        """,
    )

    add_heading(doc, "六、实验结果与分析", 1)
    add_heading(doc, "6.1 Toy 数据集流程验证结果", 2)
    add_table(
        doc,
        ["数据集", "Epoch", "Accuracy", "Precision", "Recall", "F1", "说明"],
        [["Toy Val", "8", "1.000", "1.000", "1.000", "1.000", "用于验证工程流程，不作为真实道路最终性能"]],
    )
    add_paragraph(
        doc,
        "在 toy 数据集上，模型训练 8 个 epoch 后验证集 accuracy、precision、recall 和 F1 均达到 1.000。该结果说明当前模型和训练代码可以正常收敛，数据加载、指标计算、模型保存和 TensorBoard 记录流程均已跑通。但 toy 数据集由脚本合成，样本规律较简单，不能代表真实道路环境下的泛化性能。",
    )
    add_paragraph(
        doc,
        "TensorBoard 中记录了 train 与 val 两组指标，每组包含 loss、accuracy、precision、recall 和 F1。训练过程中 loss 整体下降，F1 和 accuracy 逐步上升，说明模型在该验证数据上学习到了区分车道内外的视觉模式。",
    )
    add_figure(
        doc,
        REPORT_ASSETS / "tensorboard_tusimple_curves.png",
        "图 1  TuSimple 真实训练过程 TensorBoard 曲线",
        width=6.6,
    )

    add_heading(doc, "6.2 TuSimple 真实数据训练与测试结果", 2)
    add_paragraph(
        doc,
        "TuSimple 数据集下载并解压后，实验使用 label_data_0313.json 生成训练集，使用 label_data_0531.json 生成验证集，使用 label_data_0601.json 生成独立测试集。转换后训练集包含 2857 个 in_lane 样本和 5642 个 out_lane 样本；验证集包含 358 个 in_lane 样本和 702 个 out_lane 样本；测试集包含 410 个 in_lane 样本和 814 个 out_lane 样本。",
    )
    add_table(
        doc,
        ["阶段", "in_lane", "out_lane", "总样本", "用途"],
        [
            ["Train", "2857", "5642", "8499", "模型参数学习"],
            ["Val", "358", "702", "1060", "选择 best.pt"],
            ["Test", "410", "814", "1224", "最终独立评估"],
        ],
    )
    add_paragraph(
        doc,
        "模型共训练 20 个 epoch，验证集 F1 最高出现在第 18 个 epoch，因此将 outputs/tusimple_run/best.pt 固定为正式模型。该模型在验证集上的 accuracy 为 0.9085，precision 为 0.8254，recall 为 0.9246，F1 为 0.8722。",
    )
    add_table(
        doc,
        ["模型", "Epoch", "Loss", "Accuracy", "Precision", "Recall", "F1"],
        [["best.pt", "18", "0.2320", "0.9085", "0.8254", "0.9246", "0.8722"]],
    )
    add_paragraph(
        doc,
        "在独立测试集上，正式模型取得 accuracy 0.9665、precision 0.9301、recall 0.9732、F1 0.9511、loss 0.0912。该结果说明模型在未参与训练和模型选择的新 ROI 样本上仍能较稳定地区分车道内与车道外。",
    )
    add_table(
        doc,
        ["Split", "Loss", "Accuracy", "Precision", "Recall", "F1", "Samples"],
        [["Test", "0.0912", "0.9665", "0.9301", "0.9732", "0.9511", "1224"]],
    )
    add_table(
        doc,
        ["指标", "数量", "含义"],
        [
            ["TP", "399", "真实为 in_lane，预测为 in_lane"],
            ["TN", "784", "真实为 out_lane，预测为 out_lane"],
            ["FP", "30", "真实为 out_lane，误判为 in_lane"],
            ["FN", "11", "真实为 in_lane，误判为 out_lane"],
        ],
    )
    add_paragraph(
        doc,
        "从混淆矩阵可以看出，模型错误主要集中在 FP，即部分 out_lane 样本被误判为 in_lane。后续可以通过增加困难负样本、调整分类阈值或优化 ROI 偏移策略继续降低该类误判。本实验当前阶段的 F1 已达到 0.9511，满足课程作业对简单模型验证的要求。",
    )

    add_heading(doc, "6.3 ONNX 与 ONNX Runtime 验证", 2)
    add_table(
        doc,
        ["样本", "预测类别", "probability_in_lane", "logit", "推理后端"],
        [
            ["test/in_lane/..._center.png", "in_lane", "0.9432", "2.8098", "CPUExecutionProvider"],
            ["test/out_lane/..._right.png", "out_lane", "0.0000", "-27.0708", "CPUExecutionProvider"],
        ],
    )
    add_paragraph(
        doc,
        "ONNX Runtime 推理结果表明，导出的 lane_binary_classifier.onnx 可以脱离 PyTorch 独立运行。真实 test 集 in_lane 样本的 in_lane 概率为 0.9432，高于 0.5 阈值，因此预测为车道内；真实 test 集 out_lane 样本的 in_lane 概率接近 0，低于阈值，因此预测为车道外。这证明模型已完成从训练框架到部署推理格式的闭环验证。",
    )
    add_figure(
        doc,
        REPORT_ASSETS / "onnx_runtime_tusimple_examples.png",
        "图 2  ONNX Runtime 在真实 test ROI 上的推理结果展示",
        width=6.6,
    )
    add_code(
        doc,
        """
session = ort.InferenceSession(args.onnx_model, providers=[provider])
input_name = session.get_inputs()[0].name
output_name = session.get_outputs()[0].name
image = preprocess(args.image, image_config)
output = session.run([output_name], {input_name: image})[0]

logit = float(np.asarray(output).reshape(-1)[0])
probability = float(sigmoid(np.array([logit], dtype=np.float32))[0])
target = 1 if probability >= args.threshold else 0
        """,
    )

    add_heading(doc, "6.4 误差样例说明", 2)
    add_paragraph(
        doc,
        "在 ONNX Runtime 单图测试中，也发现一张 out_lane 样本被误判为 in_lane，probability_in_lane 为 0.7757。这与测试集混淆矩阵中 FP=30 的现象一致，说明模型对部分靠近车道边界、纹理与车道中心相似的负样本仍存在误判风险。该问题可作为后续优化方向，而不影响当前模型作为课程作业基础方案的有效性。",
    )
    add_table(
        doc,
        ["样本", "真实类别", "预测类别", "probability_in_lane", "说明"],
        [["test/out_lane/..._right.png", "out_lane", "in_lane", "0.7757", "典型 FP 误判，可用于误差分析"]],
    )

    add_heading(doc, "七、问题记录与解决方案", 1)
    add_table(
        doc,
        ["问题", "原因", "解决方法"],
        [
            ["pip 命令误连导致 .pip is not a valid editable requirement", "两条 pip 命令被连续输入成一条命令", "分开执行 pip install -r requirements.txt 和 pip install -e ."],
            ["TensorBoard 提示未安装 TensorFlow", "TensorBoard 最早属于 TensorFlow 工具链，会检测 TensorFlow", "本实验使用 PyTorch 写入 TensorBoard，提示不影响曲线显示"],
            ["ONNX 导出缺少 onnxscript", "新版 PyTorch ONNX 导出器需要额外依赖", "补充 onnx、onnxscript、onnxruntime 到 requirements.txt"],
            ["ONNX 导出在 Windows 中文终端出现编码问题", "新版导出器打印特殊符号时 GBK 控制台无法编码", "在 export_onnx.py 中使用 dynamo=False，切换到更稳定的 legacy 导出路径"],
            ["GitHub push 被 workflow 权限拒绝", "Personal Access Token 缺少 workflow scope，不能上传 .github/workflows 文件", "移除非必要 CI 文件后成功上传项目"],
        ],
    )

    add_heading(doc, "八、方案优缺点分析", 1)
    add_heading(doc, "8.1 优点", 2)
    for item in [
        "模型结构轻量，参数量较小，训练和推理速度快，适合作业和快速原型验证。",
        "ROI 输入减少了整图中的无关背景，使模型更关注车辆前方道路区域。",
        "训练流程完整，包含数据加载、增强、指标记录、模型保存和 TensorBoard 可视化。",
        "支持 ONNX 导出和 ONNX Runtime 推理，具备初步部署验证能力。",
        "数据转换脚本可将 TuSimple 原始车道线标注转换为二分类样本，便于从开源数据集开展正式实验。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8.2 不足", 2)
    for item in [
        "二分类模型只能输出车道内或车道外，无法给出精确车道线位置和可行驶区域边界。",
        "ROI 样本质量依赖裁剪策略，若车道中心估计不准确，会影响标签质量。",
        "toy 数据集与真实道路场景差异较大，只能验证流程，不能代表最终性能。",
        "当前模型未利用连续帧时序信息，遇到遮挡、模糊、弯道或强光变化时可能不够稳定。",
        "尚未统计 ONNX Runtime 的推理耗时，部署性能分析仍需补充。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "九、总结与下一阶段工作", 1)
    add_paragraph(
        doc,
        "本实验完成了车道区域二分类模型的设计与工程实现，并已基于 TuSimple 开源车道线数据集完成正式训练、验证和测试。实验流程包括 TuSimple JSON 标注解析、ROI 二分类样本构建、轻量 CNN 训练、TensorBoard 记录、PyTorch checkpoint 保存、ONNX 导出和 ONNX Runtime 独立推理。最终模型在独立测试集上取得 accuracy 0.9665、precision 0.9301、recall 0.9732、F1 0.9511，说明该方案能够较稳定地完成车道内外 ROI 判断。",
    )
    add_paragraph(
        doc,
        "下一阶段工作不再是接入 TuSimple，而是在当前真实数据实验基础上继续优化模型鲁棒性和部署表现。可进一步补充更多误差样例截图，重点分析 FP 样本，即真实为 out_lane 但被误判为 in_lane 的情况；尝试调整分类阈值以降低误判风险；比较不同 ROI 尺寸、y_ref 参考位置和负样本偏移距离对指标的影响；增加阴影、模糊、雨雾、强光等增强策略；尝试 MobileNetV3、ResNet18 等轻量骨干网络；并统计 ONNX Runtime 推理耗时，为后续边缘部署提供依据。",
    )

    add_heading(doc, "十、附录：正式 TuSimple 训练命令", 1)
    add_code(
        doc,
        """
python scripts\\prepare_tusimple_binary.py --tusimple-root D:\\datasets\\tusimple\\train_set --label-file D:\\datasets\\tusimple\\train_set\\label_data_0313.json --split train --output data\\tusimple_binary

python scripts\\prepare_tusimple_binary.py --tusimple-root D:\\datasets\\tusimple\\train_set --label-file D:\\datasets\\tusimple\\train_set\\label_data_0531.json --split val --output data\\tusimple_binary

python -m lane_binary_classifier.train --data-dir data\\tusimple_binary --output-dir outputs\\tusimple_run --epochs 20 --batch-size 32 --tensorboard

python -m lane_binary_classifier.export_onnx --checkpoint outputs\\tusimple_run\\best.pt --output outputs\\tusimple_run\\lane_binary_classifier.onnx

python -m lane_binary_classifier.evaluate --checkpoint outputs\\tusimple_run\\best.pt --data-dir data\\tusimple_binary --split test --output outputs\\tusimple_run\\test_metrics.json

python -m lane_binary_classifier.predict_onnx --onnx-model outputs\\tusimple_run\\lane_binary_classifier.onnx --image data\\tusimple_binary\\test\\in_lane\\clips_0601_1494452385593783358_20_center.png
        """,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
